"""Professional Word export for the source-backed executive audit report."""

from io import BytesIO

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLUE = "0B2545"
ACCENT = "0E7490"
LIGHT = "E8EEF5"
GRAY = "F2F4F7"
MUTED = "5B6573"
RED = "9B1C1C"
GOLD = "7A5A00"
CONTENT_DXA = 9360


def _font(run, size=10, color=BLUE, bold=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def _shade(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shade = properties.find(qn("w:shd"))
    if shade is None:
        shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), fill)
    if shade.getparent() is None:
        properties.append(shade)


def _cell_margins(cell, top=80, start=120, bottom=80, end=120):
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        if node.getparent() is None:
            margins.append(node)


def _table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    properties = table._tbl.tblPr
    width = properties.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
    width.set(qn("w:w"), str(sum(widths)))
    width.set(qn("w:type"), "dxa")
    if width.getparent() is None:
        properties.append(width)
    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    if indent.getparent() is None:
        properties.append(indent)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)
    for row in table.rows:
        for cell, value in zip(row.cells, widths):
            cell.width = Inches(value / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_width = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_width.set(qn("w:w"), str(value))
            tc_width.set(qn("w:type"), "dxa")
            _cell_margins(cell)


def _set_cell(cell, value, *, bold=False, color=BLUE, size=9, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    _font(paragraph.add_run(str(value)), size=size, color=color, bold=bold)


def _heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run(text)
    return paragraph


def _format_number(value):
    return f"{int(value or 0):,}".replace(",", ".")


def _schedule(rule):
    if not rule["is_active"]:
        return "Pasif"
    minutes = rule.get("schedule_interval_minutes") if rule.get("schedule_enabled") else None
    if not minutes:
        return "Manuel"
    if minutes % 1440 == 0:
        return f"{minutes // 1440} günde bir"
    if minutes % 60 == 0:
        return f"{minutes // 60} saatte bir"
    return f"{minutes} dakikada bir"


def _priority_label(value):
    return {"immediate": "ACİL", "high": "YÜKSEK", "medium": "ORTA",
            "low": "DÜŞÜK", "normal": "NORMAL"}.get(value, str(value).upper())


def build_executive_word_report(report):
    """Return a deterministic DOCX representation of an executive report payload."""
    document = Document()
    section = document.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.bottom_margin = Inches(0.8)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name, normal.font.size, normal.font.color.rgb = "Calibri", Pt(10), RGBColor.from_string(BLUE)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, before, after in (("Heading 1", 16, 16, 8), ("Heading 2", 13, 12, 6),
                                      ("Heading 3", 12, 8, 4)):
        style = document.styles[name]
        style.font.name, style.font.size = "Calibri", Pt(size)
        style.font.bold, style.font.color.rgb = True, RGBColor.from_string(ACCENT if name != "Heading 3" else BLUE)
        style.paragraph_format.space_before, style.paragraph_format.space_after = Pt(before), Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _font(header.add_run("AUDITAI  |  YÖNETİCİ DENETİM RAPORU"), 8.5, MUTED, True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(footer.add_run("Gizlilik: Kurum içi kullanım  •  AuditAI tarafından üretilmiştir"), 8, MUTED)

    kicker = document.add_paragraph()
    kicker.paragraph_format.space_after = Pt(5)
    _font(kicker.add_run("SÜREKLİ DENETİM VE KONTROL ANALİTİĞİ"), 9, ACCENT, True)
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    _font(title.add_run("Yönetici Denetim Raporu"), 24, BLUE, True)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    _font(subtitle.add_run("Kontrol performansı, bulgu görünümü ve çözüm öncelikleri"), 12, MUTED)

    metadata = document.add_table(rows=2, cols=2)
    metadata.style = "Table Grid"
    _table_geometry(metadata, [1800, 7560])
    _set_cell(metadata.cell(0, 0), "Rapor dönemi", bold=True, color=MUTED)
    _set_cell(metadata.cell(0, 1), f"{report['period']['from'][:10]} – {report['period']['to'][:10]}")
    _set_cell(metadata.cell(1, 0), "Üretim zamanı", bold=True, color=MUTED)
    _set_cell(metadata.cell(1, 1), report["generated_at"].replace("T", " ")[:19])

    _heading(document, "1. Yönetici özeti")
    kpi = report["kpis"]
    metrics = [
        ("Aktif kontrol", _format_number(kpi["active_rules"])),
        ("Çalıştırma", _format_number(kpi["execution_count"])),
        ("Başarı oranı", f"%{kpi['execution_success_rate'] * 100:.1f}"),
        ("Açık bulgu", _format_number(kpi["open_findings"])),
        ("Kritik açık", _format_number(kpi["critical_open_findings"])),
    ]
    strip = document.add_table(rows=2, cols=len(metrics))
    _table_geometry(strip, [CONTENT_DXA // len(metrics)] * len(metrics))
    for index, (label, value) in enumerate(metrics):
        _shade(strip.cell(0, index), BLUE)
        _shade(strip.cell(1, index), LIGHT)
        _set_cell(strip.cell(0, index), label, bold=True, color="FFFFFF", size=8.5,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell(strip.cell(1, index), value, bold=True, color=BLUE, size=15,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

    narrative = document.add_paragraph()
    narrative.paragraph_format.space_before = Pt(10)
    narrative.paragraph_format.space_after = Pt(8)
    _font(narrative.add_run("Yönetim değerlendirmesi: "), 10.5, BLUE, True)
    urgent = sum(rule["priority"] in {"immediate", "high"} for rule in report["rules"])
    _font(narrative.add_run(
        f"{urgent} kontrol öncelikli aksiyon gerektiriyor. {kpi['scanned_records']:,} kayıt tarandı, "
        f"{kpi['matched_records']:,} istisna eşleşmesi ve {kpi['finding_count']} bulgu üretildi. "
        f"Ortalama çözüm süresi {kpi['average_resolution_hours']:.1f} saattir."), 10.5)

    _heading(document, "2. Bulgular ve çözüm öncelikleri")
    priorities = document.add_table(rows=1, cols=4)
    priorities.style = "Table Grid"
    _table_geometry(priorities, [3000, 1600, 1560, 3200])
    for cell, value in zip(priorities.rows[0].cells, ("Önem seviyesi", "Bulgu", "Açık", "Yönetim yaklaşımı")):
        _shade(cell, GRAY); _set_cell(cell, value, bold=True)
    guidance = {"critical": "Derhal değerlendir", "high": "Kısa vadeli aksiyon",
                "medium": "Planlı iyileştirme", "low": "Rutin takip"}
    for severity in ("critical", "high", "medium", "low"):
        cells = priorities.add_row().cells
        total = report["findings_by_severity"].get(severity, 0)
        open_count = sum(rule["open_findings"] for rule in report["rules"] if rule["severity"] == severity)
        color = RED if severity == "critical" else GOLD if severity in {"high", "medium"} else BLUE
        for cell, value, align in zip(cells, (severity.upper(), total, open_count, guidance[severity]),
                                      (WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
                                       WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT)):
            _set_cell(cell, value, bold=cell is cells[0], color=color if cell is cells[0] else BLUE, align=align)
    _table_geometry(priorities, [3000, 1600, 1560, 3200])

    _heading(document, "3. Kaynak bazında bulgu yoğunluğu")
    sources = document.add_table(rows=1, cols=3)
    sources.style = "Table Grid"
    _table_geometry(sources, [5000, 1800, 2560])
    for cell, value in zip(sources.rows[0].cells, ("Veri kaynağı", "Bulgu sayısı", "Toplam içindeki pay")):
        _shade(cell, GRAY); _set_cell(cell, value, bold=True)
    finding_total = max(1, kpi["finding_count"])
    source_rows = list(report["findings_by_source"])
    displayed_sources = source_rows[:8]
    if len(source_rows) > 8:
        displayed_sources.append({"name": f"Diğer {len(source_rows) - 8} kaynak",
                                  "count": sum(item["count"] for item in source_rows[8:])})
    for source in displayed_sources:
        cells = sources.add_row().cells
        for cell, value, align in zip(cells, (source["name"], source["count"],
                                              f"%{source['count'] / finding_total * 100:.1f}"),
                                      (WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
                                       WD_ALIGN_PARAGRAPH.CENTER)):
            _set_cell(cell, value, align=align)
    if len(sources.rows) == 1:
        cells = sources.add_row().cells
        _set_cell(cells[0], "Seçilen dönemde bulgu oluşmadı.", color=MUTED)
        cells[0].merge(cells[2])
    _table_geometry(sources, [5000, 1800, 2560])

    document.add_page_break()
    _heading(document, "4. Kural kapsamı ve performans envanteri")
    intro = document.add_paragraph("Aşağıdaki envanter her kontrolün veri kapsamını, çalışma sıklığını, performansını ve bulgu önceliğini gösterir.")
    intro.paragraph_format.space_after = Pt(8)
    inventory = document.add_table(rows=1, cols=6)
    inventory.style = "Table Grid"
    widths = [2500, 2100, 1200, 1200, 1060, 1300]
    _table_geometry(inventory, widths)
    headers = ("Kural / alan", "Kaynaklar", "Sıklık", "Çalışma", "Açık", "Öncelik")
    for cell, value in zip(inventory.rows[0].cells, headers):
        _shade(cell, BLUE); _set_cell(cell, value, bold=True, color="FFFFFF", size=8.5)
    for rule in report["rules"]:
        cells = inventory.add_row().cells
        values = (f"{rule['name']}\n{rule['audit_area']}", ", ".join(rule["sources"]), _schedule(rule),
                  f"{rule['execution_count']} kez\n{rule['average_duration_seconds']:.2f} sn",
                  rule["open_findings"], _priority_label(rule["priority"]))
        for index, (cell, value) in enumerate(zip(cells, values)):
            _set_cell(cell, value, bold=index in {0, 5}, color=RED if rule["priority"] == "immediate" and index == 5 else BLUE,
                      size=8.2, align=WD_ALIGN_PARAGRAPH.CENTER if index >= 2 else WD_ALIGN_PARAGRAPH.LEFT)
    _table_geometry(inventory, widths)

    _heading(document, "5. Metrik tanımları ve kullanım notu")
    notes = [
        "Çalıştırma başarı oranı, seçilen dönemde tamamlanan çalışmaların tüm çalışmalara oranıdır.",
        "Eşleşme oranı, kuralların eşleştirdiği kayıtların taranan kayıtlara oranıdır; doğrudan hata oranı olarak yorumlanmamalıdır.",
        "Çözüm süresi, çözülmüş bulgularda oluşturulma ve son güncellenme zamanı arasındaki ortalamadır.",
        "Bu rapor, rapor üretildiği anda AuditAI veritabanındaki kayıtlardan hazırlanmıştır.",
    ]
    for note in notes:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(5)
        paragraph.add_run(note)

    _heading(document, "6. Yönetim aksiyon planı")
    action_table = document.add_table(rows=1, cols=3)
    action_table.style = "Table Grid"
    action_widths = [3600, 2700, 3060]
    _table_geometry(action_table, action_widths)
    for cell, value in zip(action_table.rows[0].cells, ("Öncelikli kontrol", "Kaynak", "Önerilen aksiyon")):
        _shade(cell, BLUE); _set_cell(cell, value, bold=True, color="FFFFFF", size=8.5)
    action_copy = {"immediate": "Derhal incele, sorumlu ata ve kapanış tarihi belirle.",
                   "high": "Kısa vadeli düzeltici aksiyon planı oluştur.",
                   "medium": "Planlı iyileştirme takvimine dahil et.",
                   "low": "Rutin izleme kapsamında takip et."}
    action_rules = sorted((rule for rule in report["rules"] if rule["open_findings"]),
                          key=lambda rule: ({"immediate": 0, "high": 1, "medium": 2, "low": 3,
                                             "normal": 4}.get(rule["priority"], 5),
                                            -rule["open_findings"], rule["name"]))[:8]
    for rule in action_rules:
        cells = action_table.add_row().cells
        values = (f"{rule['name']} ({rule['open_findings']} açık)", ", ".join(rule["sources"]),
                  action_copy.get(rule["priority"], "Yönetici değerlendirmesine al."))
        for cell, value in zip(cells, values):
            _set_cell(cell, value, size=8.5)
    if not action_rules:
        cells = action_table.add_row().cells
        _set_cell(cells[0], "Açık bulgu bulunmadığı için yönetim aksiyonu gerekmiyor.", color=MUTED)
        cells[0].merge(cells[2])
    _table_geometry(action_table, action_widths)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
